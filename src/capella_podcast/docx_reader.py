"""Read generated/edited DOCX files back from disk.

Regeneration must always re-parse the artifact the user may have edited —
never cached text — so report edits flow into scripts and script edits flow
into audio.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx

HOST_PREFIX = re.compile(r"^\s*(HOST\s+[A-Z])\s*[:\-]\s*(.*)$", re.IGNORECASE)


def read_docx_text(path: Path) -> str:
    """All paragraph text (including hyperlink text), one line per paragraph."""
    doc = docx.Document(str(path))
    lines = [p.text.strip() for p in doc.paragraphs]
    return "\n".join(ln for ln in lines if ln)


def read_script_turns(path: Path) -> list[tuple[str, str]]:
    """Parse speaker turns from a script DOCX.

    Returns [("HOST A", "spoken text"), ...]. Paragraphs without a HOST
    prefix continue the previous speaker's turn (a user edit may split a
    turn across paragraphs); leading non-host paragraphs (title etc.) are
    skipped.
    """
    doc = docx.Document(str(path))
    turns: list[tuple[str, str]] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = HOST_PREFIX.match(text)
        if m:
            host = m.group(1).upper().replace("  ", " ")
            spoken = m.group(2).strip()
            turns.append((host, spoken))
        elif turns:
            host, spoken = turns[-1]
            turns[-1] = (host, f"{spoken} {text}".strip())
    return [(h, t) for h, t in turns if t]
