"""Styled DOCX rendering matching the example Capella summary reports.

Styling constants (verified against the example files' OOXML):
- US Letter (12240 x 15840 DXA), 1" margins, header/footer at 720 DXA.
- Brand teal 0F4761: title 24pt, subtitle 16pt, section headings 14pt bold.
- Body 12pt. Font Aptos, fallback Arial when Aptos is not installed.
- Header: Capella logo top-left (inline, ~1.86" wide), course number at a
  right tab stop, thin 808080 rule below.
- Footer: right-aligned PAGE field with the same thin rule above.
- Bullets use real Word list numbering (List Paragraph + numPr), with a bold
  lead-in phrase; resource names become real hyperlinks.
"""

from __future__ import annotations

import sys
from dataclasses import dataclass, field as dc_field
from pathlib import Path

import docx
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from .config import AppConfig

RULE_COLOR = "808080"
LOGO_WIDTH = Inches(1.86)
RIGHT_TAB = Twips(9360)  # right margin edge for 6.5" text width


@dataclass
class Bullet:
    lead_in: str
    text: str = ""
    url: str | None = None


@dataclass
class Section:
    heading: str
    paragraphs: list[str] = dc_field(default_factory=list)
    bullets: list[Bullet] = dc_field(default_factory=list)


def find_logo(cfg: AppConfig) -> Path | None:
    """Locate the Capella logo by case-insensitive *capella*logo*.png."""
    override = cfg.resolve(cfg.report.logo_path)
    if override and override.is_file():
        return override
    assets = cfg.resolve(cfg.report.assets_dir)
    if assets and assets.is_dir():
        for p in sorted(assets.iterdir()):
            n = p.name.lower()
            if p.is_file() and n.endswith(".png") and "capella" in n and "logo" in n:
                return p
    return None


def _font_installed(name: str) -> bool:
    if sys.platform == "win32":
        import winreg

        for root in (winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER):
            try:
                key = winreg.OpenKey(root, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts")
            except OSError:
                continue
            try:
                for i in range(winreg.QueryInfoKey(key)[1]):
                    if winreg.EnumValue(key, i)[0].lower().startswith(name.lower()):
                        return True
            finally:
                key.Close()
        return False
    candidates = ["/Library/Fonts", "/System/Library/Fonts", "~/Library/Fonts",
                  "/usr/share/fonts", "~/.local/share/fonts", "~/.fonts"]
    return any(
        True
        for d in candidates
        for p in Path(d).expanduser().rglob(f"{name}*")
        if Path(d).expanduser().is_dir()
    )


def pick_font(cfg: AppConfig) -> str:
    if _font_installed(cfg.report.font):
        return cfg.report.font
    print(
        f"NOTE: font {cfg.report.font!r} not installed; using {cfg.report.font_fallback!r}.",
        file=sys.stderr,
    )
    return cfg.report.font_fallback


class SummaryDocBuilder:
    def __init__(self, cfg: AppConfig, course_number: str):
        self.cfg = cfg
        self.font = pick_font(cfg)
        self.color = RGBColor.from_string(cfg.report.brand_color)
        self.doc: DocumentObject = docx.Document()
        self._setup_page()
        self._setup_header_footer(course_number)

    # -- page / chrome -------------------------------------------------

    def _setup_page(self) -> None:
        sect = self.doc.sections[0]
        sect.page_width = Twips(12240)
        sect.page_height = Twips(15840)
        for side in ("top", "bottom", "left", "right"):
            setattr(sect, f"{side}_margin", Inches(1))
        sect.header_distance = Twips(720)
        sect.footer_distance = Twips(720)
        style = self.doc.styles["Normal"]
        style.font.name = self.font
        style.font.size = Pt(12)

    def _para_border(self, paragraph, edge: str) -> None:
        """Thin horizontal rule on one edge of a paragraph (matches examples)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            pPr.append(pBdr)
        el = pBdr.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "3")
        el.set(qn("w:color"), RULE_COLOR)
        pBdr.append(el)

    def _setup_header_footer(self, course_number: str) -> None:
        header = self.doc.sections[0].header
        hp = header.paragraphs[0]
        hp.text = ""
        self._para_border(hp, "bottom")
        pPr = hp._p.get_or_add_pPr()
        tabs = pPr.makeelement(qn("w:tabs"), {})
        tab = tabs.makeelement(qn("w:tab"), {})
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9360")
        tabs.append(tab)
        pPr.append(tabs)

        logo = find_logo(self.cfg)
        if logo is not None:
            hp.add_run().add_picture(str(logo), width=LOGO_WIDTH)
        else:
            print(
                "WARNING: no Capella logo found in assets/ (*capella*logo*.png); "
                "rendering header without it.",
                file=sys.stderr,
            )
        hp.add_run("\t")
        run = hp.add_run(course_number)
        run.font.name = self.font
        run.font.size = Pt(12)

        footer = self.doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.text = ""
        self._para_border(fp, "top")
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._add_page_field(fp)

    def _add_page_field(self, paragraph) -> None:
        run = paragraph.add_run()
        run.font.name = self.font
        run.font.size = Pt(12)
        for el_name, attrs, text in (
            ("w:fldChar", {"w:fldCharType": "begin"}, None),
            ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
            ("w:fldChar", {"w:fldCharType": "end"}, None),
        ):
            el = run._r.makeelement(qn(el_name), {})
            for k, v in attrs.items():
                el.set(qn(k), v)
            if text:
                el.text = text
            run._r.append(el)

    # -- content -------------------------------------------------------

    def _styled_run(self, paragraph, text: str, size: float, bold: bool = False,
                    color: RGBColor | None = None):
        run = paragraph.add_run(text)
        run.font.name = self.font
        run.font.size = Pt(size)
        run.bold = bold or None
        if color is not None:
            run.font.color.rgb = color
        return run

    def add_title(self, text: str) -> None:
        p = self.doc.add_paragraph()
        self._styled_run(p, text, 24, color=self.color)

    def add_subtitle(self, text: str) -> None:
        p = self.doc.add_paragraph()
        self._styled_run(p, text, 16, color=self.color)

    def add_heading(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        self._styled_run(p, text, 14, bold=True, color=self.color)

    def add_body(self, text: str) -> None:
        p = self.doc.add_paragraph()
        self._styled_run(p, text, 12)

    def add_bullet(self, bullet: Bullet) -> None:
        p = self.doc.add_paragraph(style="List Bullet")
        lead = bullet.lead_in.rstrip(": ")
        if bullet.url:
            self._add_hyperlink_run(p, f"{lead}", bullet.url, bold=True)
        else:
            self._styled_run(p, lead, 12, bold=True)
        if bullet.text:
            self._styled_run(p, ": " if not bullet.url else ": ", 12, bold=True)
            self._styled_run(p, bullet.text, 12)

    def _add_hyperlink_run(self, paragraph, text: str, url: str, bold: bool = False) -> None:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {})
        hyperlink.set(qn("r:id"), r_id)
        r = hyperlink.makeelement(qn("w:r"), {})
        rPr = r.makeelement(qn("w:rPr"), {})
        fonts = rPr.makeelement(qn("w:rFonts"), {})
        fonts.set(qn("w:ascii"), self.font)
        fonts.set(qn("w:hAnsi"), self.font)
        rPr.append(fonts)
        if bold:
            rPr.append(rPr.makeelement(qn("w:b"), {}))
        sz = rPr.makeelement(qn("w:sz"), {})
        sz.set(qn("w:val"), "24")  # half-points: 12pt
        rPr.append(sz)
        color = rPr.makeelement(qn("w:color"), {})
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        u = rPr.makeelement(qn("w:u"), {})
        u.set(qn("w:val"), "single")
        rPr.append(u)
        r.append(rPr)
        t = r.makeelement(qn("w:t"), {})
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        hyperlink.append(r)
        paragraph._p.append(hyperlink)

    def add_section(self, section: Section) -> None:
        self.add_heading(section.heading)
        for para in section.paragraphs:
            self.add_body(para)
        for b in section.bullets:
            self.add_bullet(b)

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def render_summary_docx(
    cfg: AppConfig,
    path: Path,
    title: str,
    subtitle: str,
    course_number: str,
    sections: list[Section],
) -> None:
    b = SummaryDocBuilder(cfg, course_number)
    b.add_title(title)
    b.add_subtitle(subtitle)
    for s in sections:
        b.add_section(s)
    b.save(path)
