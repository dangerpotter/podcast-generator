"""Validation for every generated DOCX (spec: validate before moving on)."""

from __future__ import annotations

from pathlib import Path

import docx
from docx.oxml.ns import qn

PLACEHOLDER = "Remove or Replace"
BULLET_CHARS = ("•", "- ", "* ", "·")


def _has_numbering(paragraph) -> bool:
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return True
    style = paragraph.style
    while style is not None:
        if style.name and style.name.startswith(("List Bullet", "List Number", "List Paragraph")):
            return True
        style = style.base_style
    return False


def validate_docx(
    path: Path,
    expected_headings: list[str],
    title: str | None = None,
    min_hyperlinks: int = 0,
    expect_logo: bool = False,
    course_number: str | None = None,
) -> list[str]:
    """Return a list of problems (empty = valid)."""
    problems: list[str] = []
    try:
        doc = docx.Document(str(path))
    except Exception as e:
        return [f"file does not open as DOCX: {e}"]

    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    joined = "\n".join(texts)

    if title and not any(t.startswith(title) for t in texts):
        problems.append(f"missing title {title!r}")
    for h in expected_headings:
        if not any(t == h or t.startswith(h) for t in texts):
            problems.append(f"missing section heading {h!r}")
    if PLACEHOLDER.lower() in joined.lower():
        problems.append(f"placeholder text {PLACEHOLDER!r} present in body")

    # bullets: real Word numbering, no literal bullet characters
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and any(t.startswith(c) for c in BULLET_CHARS):
            problems.append(f"literal bullet character in: {t[:60]!r}")
    n_bulleted = sum(1 for p in doc.paragraphs if _has_numbering(p) and p.text.strip())
    if n_bulleted == 0:
        problems.append("no real Word list numbering found")

    n_links = sum(1 for r in doc.part.rels.values() if r.reltype.endswith("/hyperlink"))
    if n_links < min_hyperlinks:
        problems.append(f"expected >= {min_hyperlinks} hyperlinks, found {n_links}")

    sect = doc.sections[0]
    if (sect.page_width, sect.page_height) != (12240 * 635, 15840 * 635):
        problems.append(f"page size is not US Letter ({sect.page_width}x{sect.page_height} EMU)")
    if any(
        getattr(sect, f"{side}_margin") != 914400 for side in ("top", "bottom", "left", "right")
    ):
        problems.append("margins are not 1 inch")

    header_text = "\n".join(p.text for p in sect.header.paragraphs)
    if PLACEHOLDER.lower() in header_text.lower():
        problems.append(f"placeholder text {PLACEHOLDER!r} present in header")
    if course_number and course_number not in header_text:
        problems.append(f"course number {course_number!r} missing from header")
    if expect_logo:
        has_image = bool(sect.header._element.findall(".//" + qn("w:drawing")))
        if not has_image:
            problems.append("logo image missing from header")

    footer_xml = sect.footer._element.xml
    if "PAGE" not in footer_xml:
        problems.append("footer page number field missing")
    if "pBdr" not in footer_xml:
        problems.append("footer rule (border) missing")
    return problems
