"""Compare a generated summary DOCX's layout fingerprint to the example reports.

The examples are layout references: US Letter, 1" margins, teal 0F4761 title
and headings, 24/16/14pt heading scale, 12pt body, logo + course number in the
header, ruled page-number footer, real Word list numbering with bold lead-ins.

Usage:
    python tests/compare_layout.py <generated.docx> [more.docx ...]
"""

from __future__ import annotations

import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

BRAND = "0F4761"


def fingerprint(path: Path) -> dict:
    doc = docx.Document(str(path))
    sect = doc.sections[0]
    fp = {
        "page_letter": (sect.page_width, sect.page_height) == (12240 * 635, 15840 * 635),
        "margins_1in": all(
            getattr(sect, f"{s}_margin") == 914400 for s in ("top", "bottom", "left", "right")
        ),
        "header_logo": bool(sect.header._element.findall(".//" + qn("w:drawing"))),
        "header_rule": "pBdr" in sect.header._element.xml,
        "footer_page_field": "PAGE" in sect.footer._element.xml,
        "footer_rule": "pBdr" in sect.footer._element.xml,
        "placeholder_absent": "remove or replace" not in (
            "\n".join(p.text for p in doc.paragraphs)
            + sect.header._element.xml
        ).lower(),
    }

    sizes: dict[str, float | None] = {"title": None, "subtitle": None, "heading": None, "body": None}
    colors_ok, bullets, bold_leads, literal_bullets = True, 0, 0, 0
    seen_title = seen_sub = False
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        runs = [r for r in p.runs if r.text.strip()]
        if not runs:
            continue
        sz = runs[0].font.size
        color = runs[0].font.color.rgb if runs[0].font.color and runs[0].font.color.type else None
        is_listed = (p.style.name or "").startswith("List") or p._p.find(
            qn("w:pPr")) is not None and p._p.find(qn("w:pPr")).find(qn("w:numPr")) is not None
        if not seen_title:
            sizes["title"] = sz.pt if sz else None
            colors_ok &= str(color) == BRAND
            seen_title = True
        elif not seen_sub:
            sizes["subtitle"] = sz.pt if sz else None
            colors_ok &= str(color) == BRAND
            seen_sub = True
        elif str(color) == BRAND:
            sizes["heading"] = sz.pt if sz else None
        elif is_listed:
            bullets += 1
            if runs[0].bold:
                bold_leads += 1
        elif sizes["body"] is None:
            sizes["body"] = sz.pt if sz else None
        if p.text.strip().startswith(("•", "- ", "* ")):
            literal_bullets += 1

    fp.update(
        title_24pt=sizes["title"] == 24.0,
        subtitle_16pt=sizes["subtitle"] == 16.0,
        headings_14pt=sizes["heading"] == 14.0,
        body_12pt=sizes["body"] == 12.0,
        brand_color_on_title=colors_ok,
        real_bullets=bullets > 0,
        bold_lead_ins=bullets > 0 and bold_leads == bullets,
        no_literal_bullets=literal_bullets == 0,
    )
    return fp


def main() -> None:
    failed = False
    for arg in sys.argv[1:]:
        fp = fingerprint(Path(arg))
        bad = [k for k, v in fp.items() if not v]
        print(f"\n== {arg}")
        for k, v in fp.items():
            print(f"  [{'ok' if v else 'FAIL'}] {k}")
        failed |= bool(bad)
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()
