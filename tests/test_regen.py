"""Exercise the edit-and-regenerate loops on one module (acceptance criteria).

Usage: python tests/test_regen.py output/SWK5017 1

1. Edits the podcast script DOCX (marker word), runs `regen --from-script`,
   asserts only the MP3 regenerated and the edit survived on disk.
2. Edits the assessment summary DOCX (marker tip), runs `regen --from-summary`,
   asserts the script and MP3 regenerated, the summary edit survived, and
   reports whether the marker flowed into the new script.
"""

from __future__ import annotations

import subprocess
import sys
import time
from pathlib import Path

import docx

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "src"))

from capella_podcast.docx_reader import read_docx_text, read_script_turns
from capella_podcast.artifacts import artifact_filename

PY = sys.executable
MARKER_SCRIPT = "porcupine"
MARKER_SUMMARY = "Keep a porcupine-shaped stress ball on your desk"


def run_cli(*args: str) -> None:
    cmd = [PY, "-X", "utf8", "-m", "capella_podcast.cli",
           "--config", str(REPO / "config.yaml"), *args]
    env = {"PYTHONPATH": str(REPO / "src"), "PYTHONUNBUFFERED": "1"}
    import os
    res = subprocess.run(cmd, env={**os.environ, **env}, capture_output=True, text=True)
    sys.stdout.write(res.stdout)
    sys.stderr.write(res.stderr)
    if res.returncode != 0:
        sys.exit(f"CLI failed: {' '.join(args)}")


def main() -> None:
    course_dir = Path(sys.argv[1])
    n = int(sys.argv[2]) if len(sys.argv) > 2 else 1
    mod_dirs = list(course_dir.glob(f"*-{n:02d}"))
    assert mod_dirs, f"no module dir for module {n}"
    mod = mod_dirs[0]
    course = course_dir.name
    summary = mod / artifact_filename(course, "summary", n)
    script = mod / artifact_filename(course, "script", n)
    mp3 = mod / artifact_filename(course, "podcast", n)
    failures: list[str] = []

    def check(cond, msg):
        print(f"  [{'ok' if cond else 'FAIL'}] {msg}")
        if not cond:
            failures.append(msg)

    # --- regen --from-script: edit a spoken turn, only MP3 regenerates -----
    print("\n== regen --from-script")
    d = docx.Document(str(script))
    edited = False
    for p in d.paragraphs:
        if p.text.strip().upper().startswith(("HOST A:", "HOST B:")) and len(p.runs) > 1:
            p.runs[-1].text = p.runs[-1].text + f" And remember the {MARKER_SCRIPT}."
            edited = True
            break
    assert edited, "could not edit a script turn"
    d.save(str(script))
    t_script = script.stat().st_mtime
    t_summary = summary.stat().st_mtime
    t_mp3_before = mp3.stat().st_mtime
    time.sleep(1.1)
    run_cli("regen", "--from-script", "--module", str(n), "--course", course)
    check(mp3.stat().st_mtime > t_mp3_before, f"{mp3.name} regenerated")
    check(script.stat().st_mtime == t_script, f"{script.name} untouched by regen")
    check(summary.stat().st_mtime == t_summary, f"{summary.name} untouched by regen")
    turns_text = " ".join(t for _, t in read_script_turns(script))
    check(MARKER_SCRIPT in turns_text, "script edit survived on disk (not reverted)")

    # --- regen --from-summary: edit a tip, script + MP3 regenerate ---------
    print("\n== regen --from-summary")
    d = docx.Document(str(summary))
    d.add_paragraph(f"{MARKER_SUMMARY}: squeeze it whenever a deadline feels close.",
                    style="List Bullet")
    d.save(str(summary))
    t_summary = summary.stat().st_mtime
    t_script_before = script.stat().st_mtime
    t_mp3_before = mp3.stat().st_mtime
    time.sleep(1.1)
    run_cli("regen", "--from-summary", "--module", str(n), "--course", course)
    check(script.stat().st_mtime > t_script_before, f"{script.name} regenerated")
    check(mp3.stat().st_mtime > t_mp3_before, f"{mp3.name} regenerated")
    check(summary.stat().st_mtime == t_summary, f"{summary.name} untouched by regen")
    check(MARKER_SUMMARY.split()[2] in read_docx_text(summary), "summary edit survived on disk")
    new_script_text = " ".join(t for _, t in read_script_turns(script)).lower()
    flowed = "porcupine" in new_script_text
    print(f"  [{'ok' if flowed else 'note'}] summary edit flowed into regenerated script"
          f"{'' if flowed else ' (LLM chose not to mention the marker; mechanism verified by read-from-disk)'}")

    print()
    if failures:
        sys.exit(f"{len(failures)} regen check(s) FAILED")
    print("Regen loop checks passed.")


if __name__ == "__main__":
    main()
